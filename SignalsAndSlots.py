# -*- coding: utf-8 -*-

import time
import subprocess
from PySide6.QtCore import QObject, QThread, Signal
from pandas import read_excel, ExcelFile
import docx


# Signals must inherit QObject
class MySignals(QObject):
    log_add = Signal(str)
    result_add = Signal(str)
    result_clear = Signal()
    pbar_change = Signal(int)
    btn_start_enable = Signal(bool)


# dealt with excel thread
class DealtExcelThread(QThread):
    def __init__(self, parent):
        QThread.__init__(self, parent=None)
        # Instantiate signals and connect signals to the slots
        self.signals = MySignals()
        self.signals.log_add.connect(parent.log_add)
        self.signals.result_add.connect(parent.result_add)
        self.signals.result_clear.connect(parent.result_clear)
        self.signals.pbar_change.connect(parent.pbar_change)
        self.signals.btn_start_enable.connect(parent.btn_start_enable)
        self.p = parent
        self.now = lambda: time.time()
 
    def run(self):
        try:
            filter_type = self.p.ui.comboBox_type.currentIndex()
            # filter_type 0:文本, 1:数值
            self.p.key = self.p.ui.lineEdit.text()
            self.signals.result_clear.emit()
            self.p.result = []
            self.signals.log_add.emit(f'{"*"*20}开始查找{"*"*20}')
            self.signals.btn_start_enable.emit(False)
            st = self.now()
            i,total = 0,len(self.p.fileList)
            # filter the key
            for file in self.p.fileList:
                i+=1
                try:
                    sheet_names = ExcelFile(file).sheet_names
                    self.signals.log_add.emit(f'读取 {file}...')
                    # filter the str
                    if filter_type == 0:
                        for sheet_name in sheet_names:
                            df = read_excel(file, sheet_name=sheet_name)
                            file_name = file.split("\\")[-1]
                            if True in df.columns.astype(str).str.contains(str(self.p.key)):
                                self.p.result.append((file, sheet_name))
                                self.signals.result_add.emit(f'{file_name} -> {sheet_name}.')
                                break
                            for col in df.columns:
                                if str(df[col].dtypes) == 'object':
                                    if True in df[col].str.contains(str(self.p.key)).values:
                                        self.p.result.append((file, sheet_name))
                                        self.signals.result_add.emit(f'{file_name} -> {sheet_name}.')
                                        break
                            else:
                                continue
                            break
                    # filter the num
                    elif filter_type == 1:
                        for sheet_name in sheet_names:
                            df = read_excel(file, sheet_name=sheet_name)
                            file_name = file.split("\\")[-1]
                            num = float(self.p.key)
                            if num in df.columns.values:
                                self.p.result.append((file, sheet_name))
                                self.signals.result_add.emit(f'{file_name} -> {sheet_name}.')
                                break
                            for col in df.columns:
                                if num in df[col].values:
                                    self.p.result.append((file, sheet_name))
                                    self.signals.result_add.emit(f'{file_name} -> {sheet_name}.')
                                    break
                            else:
                                continue
                            break
                except Exception as e:
                    self.signals.log_add.emit(f'{"*"*8}发生错误{"*"*8}')
                    self.signals.log_add.emit(f'错误:{str(e)},跳过本文件,继续查找...')

                self.signals.pbar_change.emit(int(i*100/total))
            # finish
            if self.p.result:
                self.signals.log_add.emit(f'{"*"*20}查找完成{"*"*20}')
                self.signals.log_add.emit('共查找{}个文件,用时{:.2f}s.'.format(len(self.p.fileList),self.now()-st))
                self.signals.log_add.emit(f'双击右侧结果可以打开(需要安装wps或者office)')
            else:
                self.signals.log_add.emit(f'{"*"*20}未查找到相应文件{"*"*20}')
                self.signals.log_add.emit('共查找{}个文件,用时{:.2f}s.'.format(len(self.p.fileList),self.now()-st))
            self.stop()

        except Exception as e:
            self.signals.log_add.emit(f'{"*"*20}发生错误,查找停止{"*"*20}')
            self.signals.log_add.emit(f'错误: {str(e)}')
            self.stop()

    def stop(self):
        self.signals.btn_start_enable.emit(True)
        self.quit()


# dealt with word thread
class DealtWordThread(QThread):
    def __init__(self, parent):
        QThread.__init__(self, parent=None)
        # Instantiate signals and connect signals to the slots
        self.signals = MySignals()
        self.signals.log_add.connect(parent.log_add)
        self.signals.result_add.connect(parent.result_add)
        self.signals.result_clear.connect(parent.result_clear)
        self.signals.pbar_change.connect(parent.pbar_change)
        self.signals.btn_start_enable.connect(parent.btn_start_enable)
        self.p = parent
        self.k = str(self.p.key)
        self.now = lambda: time.time()
 
    def run(self):
        try:
            self.p.key = self.p.ui.lineEdit.text()
            self.signals.result_clear.emit()
            self.p.result = []
            self.signals.log_add.emit(f'{"*"*20}开始查找{"*"*20}')
            self.signals.btn_start_enable.emit(False)
            st = self.now()
            # filter the key
            i,total = 0,len(self.p.fileList)
            for file in self.p.fileList:
                self.signals.log_add.emit(f'读取 {file}...')
                i+=1
                try:
                    file_name = file.split("\\")[-1]
                    self.doc = docx.Document(file)
                    for paragraph in self.doc.paragraphs:
                        if not paragraph.text.isspace() and self.k in paragraph.text:
                            self.p.result.append((file,0))
                            self.signals.result_add.emit(f'{file_name}.')
                            break
                    
                    table_content_str = ''
                    for table in self.doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                table_content_str += cell.text
                    if self.k in table_content_str:
                        self.p.result.append((file,1))
                        self.signals.result_add.emit(f'{file_name}.')
                    
                    self.signals.pbar_change.emit(int(i*100/total))
                except Exception as e:
                    self.signals.log_add.emit(f'{"*"*8}发生错误{"*"*8}')
                    self.signals.log_add.emit(f'错误:{str(e)},跳过本文件,继续查找...')

            # finish
            if self.p.result:
                self.signals.log_add.emit(f'{"*"*20}查找完成{"*"*20}')
                self.signals.log_add.emit('共查找{}个文件,用时{:.2f}s.'.format(len(self.p.fileList),self.now()-st))
                self.signals.log_add.emit(f'双击右侧结果可以打开(需要安装wps或者office)')
            else:
                self.signals.log_add.emit(f'{"*"*20}未查找到相应文件{"*"*20}')
                self.signals.log_add.emit('共查找{}个文件,用时{:.2f}s.'.format(len(self.p.fileList),self.now()-st))
            self.stop()

        except Exception as e:
            self.signals.log_add.emit(f'{"*"*20}发生错误,查找停止{"*"*20}')
            self.signals.log_add.emit(f'错误: {str(e)}')
            self.stop()

    def stop(self):
        self.signals.btn_start_enable.emit(True)
        self.quit()


# double click open file thread
class OpenFileThread(QThread):
    def __init__(self, parent,index):
        QThread.__init__(self, parent=None)
        # Instantiate signals and connect signals to the slots
        self.signals = MySignals()
        self.signals.log_add.connect(parent.log_add)
        self.p = parent
        self.index = index
 
    def run(self):
        try:
            if self.index != None:
                self.signals.log_add.emit(f'正在尝试打开 {self.p.result[self.index][0]},请稍后...')
                subprocess.call(args=f"explorer {self.p.result[self.index][0]}", shell=True)
                self.stop()

        except Exception as e:
            self.signals.log_add.emit(f'发生错误 {str(e)}')
            self.stop()

    def stop(self):
        self.quit()

